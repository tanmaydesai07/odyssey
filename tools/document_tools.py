"""
Document tools: draft generation, editing, export, authority finder, checklist.
All tools are AI-powered via llm_utils.generate() — no hardcoded stubs.
Drafts are persisted to disk so the editor and exporter can find them by ID.
"""
from smolagents.tools import Tool
from tools.llm_utils import generate, SYSTEM_PROMPTS
import json
import hashlib
import os
from datetime import datetime
from pathlib import Path

# Drafts stored here, keyed by draft_id
DRAFTS_DIR = Path(__file__).resolve().parents[1] / "drafts"
DRAFTS_DIR.mkdir(exist_ok=True)


def _draft_path(draft_id: str) -> Path:
    return DRAFTS_DIR / f"{draft_id}.json"


def _save_draft(draft_id: str, data: dict) -> None:
    with open(_draft_path(draft_id), "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def _load_draft(draft_id: str) -> dict | None:
    p = _draft_path(draft_id)
    if not p.exists():
        return None
    with open(p, "r", encoding="utf-8") as f:
        return json.load(f)


# ---------------------------------------------------------------------------
# 1. Draft Generator — AI-powered, fills template with user's actual facts
# ---------------------------------------------------------------------------

class DraftGeneratorTool(Tool):
    name = "draft_generator"
    description = (
        "Generates a complete, submission-ready draft document (FIR, consumer complaint, "
        "RTI application, labour complaint, legal notice) filled with the user's actual facts. "
        "Returns the draft text and a draft_id for later editing or export."
    )
    inputs = {
        "template_type": {
            "type": "string",
            "description": "Document type: fir | consumer_complaint | rti | labour_complaint | legal_notice",
        },
        "extracted_facts": {
            "type": "string",
            "description": (
                "User's facts as a JSON string or plain text. Include: complainant name, "
                "address, incident date, what happened, parties involved, evidence available, "
                "desired remedy."
            ),
        },
        "language": {
            "type": "string",
            "description": "Output language code: en | hi | mr | ta | te | bn | kn | ml | gu",
            "nullable": True,
        },
    }
    output_type = "string"

    # Detailed per-template system prompts
    _TEMPLATE_PROMPTS = {
        "fir": """You are a legal document drafter for Indian law. Draft a formal First Information Report (FIR) complaint letter to be submitted at a police station.

Structure:
1. To: The Station House Officer, [Police Station], [City]
2. Subject: Complaint regarding [brief issue]
3. Respectfully submitted by: [complainant details]
4. Facts of the case (numbered paragraphs): who, what, when, where, how
5. Evidence available
6. Relief sought: request to register FIR under relevant IPC/BNS sections
7. Declaration of truth
8. Date and signature line

Use formal legal language. Fill ALL fields with the provided facts. Do not leave placeholders like [NAME] — use the actual values from the facts. If a fact is missing, note it clearly.""",

        "consumer_complaint": """You are a legal document drafter for Indian law. Draft a formal consumer complaint to be filed before the District Consumer Disputes Redressal Commission under the Consumer Protection Act, 2019.

Structure:
1. Before: The District Consumer Disputes Redressal Commission, [District]
2. Complaint No.: (leave blank)
3. Complainant: [full name, address, contact]
4. Opposite Party: [seller/company name, address]
5. Subject: Complaint under Section 35 of Consumer Protection Act, 2019
6. Facts of the complaint (numbered): purchase details, defect/deficiency, attempts to resolve
7. Legal grounds: relevant sections of CPA 2019
8. Relief sought: refund/replacement/compensation amount
9. Documents annexed (list)
10. Prayer
11. Verification and declaration
12. Date and place

Fill ALL fields with provided facts. Use formal legal language.""",

        "rti": """You are a legal document drafter for Indian law. Draft a formal RTI (Right to Information) application under the Right to Information Act, 2005.

Structure:
1. To: The Public Information Officer, [Department/Ministry], [Address]
2. Subject: Application under Right to Information Act, 2005
3. Applicant details: name, address, contact
4. Information sought (numbered list — be specific and precise)
5. Period to which information relates
6. Format requested (certified copies / inspection / electronic)
7. Fee: Rs. 10/- (mention payment mode)
8. Declaration
9. Date and signature

Be specific in the information sought. Fill all fields from provided facts.""",

        "labour_complaint": """You are a legal document drafter for Indian law. Draft a formal labour complaint to be filed before the Labour Commissioner / Labour Court.

Structure:
1. To: The Labour Commissioner, [State/District]
2. Subject: Complaint regarding [wage theft / wrongful termination / harassment / etc.]
3. Complainant: [worker details, designation, employer name]
4. Respondent: [employer/company details]
5. Facts (numbered): employment details, incident, attempts to resolve
6. Applicable law: relevant sections of labour laws (Payment of Wages Act, Industrial Disputes Act, etc.)
7. Relief sought
8. Documents enclosed
9. Declaration and date

Fill all fields from provided facts.""",

        "legal_notice": """You are a legal document drafter for Indian law. Draft a formal legal notice.

Structure:
1. LEGAL NOTICE
2. From: [sender name, address, contact] — Through/By
3. To: [recipient name, address]
4. Date
5. Subject: Legal Notice for [brief issue]
6. Facts (numbered paragraphs)
7. Legal basis
8. Demand: what must be done within [X] days
9. Consequence of non-compliance
10. Without prejudice clause

Use formal legal language. Fill all fields from provided facts.""",
    }

    def forward(self, template_type: str, extracted_facts: str, language: str = "en") -> str:
        ttype = template_type.lower().strip()
        system = self._TEMPLATE_PROMPTS.get(ttype, SYSTEM_PROMPTS["draft_generator"])

        lang_note = ""
        if language and language != "en":
            lang_names = {
                "hi": "Hindi", "mr": "Marathi", "ta": "Tamil", "te": "Telugu",
                "bn": "Bengali", "kn": "Kannada", "ml": "Malayalam", "gu": "Gujarati",
            }
            lang_name = lang_names.get(language, language)
            lang_note = f"\n\nIMPORTANT: Write the entire document in {lang_name}. Keep legal section references (IPC, CPA, etc.) in English."

        prompt = (
            f"Document Type: {template_type}\n\n"
            f"User Facts:\n{extracted_facts}\n\n"
            f"Generate the complete, filled-in document ready for submission.{lang_note}"
        )

        draft_text = generate(prompt, system)

        # Persist draft to disk
        draft_id = hashlib.md5(f"{template_type}{extracted_facts}{datetime.utcnow().isoformat()}".encode()).hexdigest()[:12]
        _save_draft(draft_id, {
            "draft_id": draft_id,
            "template_type": ttype,
            "language": language or "en",
            "extracted_facts": extracted_facts,
            "draft_text": draft_text,
            "version": 1,
            "created_at": datetime.utcnow().isoformat(),
            "updated_at": datetime.utcnow().isoformat(),
        })

        return json.dumps({
            "draft_id": draft_id,
            "template_type": ttype,
            "language": language or "en",
            "draft_text": draft_text,
            "note": f"Draft saved. Use draft_id='{draft_id}' to edit or export.",
        })


# ---------------------------------------------------------------------------
# 2. Draft Editor — saves edits, can regenerate a clean version
# ---------------------------------------------------------------------------

class DraftEditorTool(Tool):
    name = "draft_editor"
    description = (
        "Edit a previously generated draft. Saves user edits and optionally regenerates "
        "a clean final version incorporating the changes."
    )
    inputs = {
        "draft_id": {
            "type": "string",
            "description": "The draft_id returned by draft_generator.",
        },
        "user_edits": {
            "type": "string",
            "description": "The user's edited version of the draft text, or instructions like 'change the date to 15 Jan 2025'.",
            "nullable": True,
        },
        "regenerate": {
            "type": "boolean",
            "description": "If true, ask the AI to produce a clean final version incorporating the edits.",
            "nullable": True,
        },
    }
    output_type = "string"

    def forward(self, draft_id: str, user_edits: str = None, regenerate: bool = False) -> str:
        draft = _load_draft(draft_id)
        if draft is None:
            return json.dumps({"error": f"Draft '{draft_id}' not found. Please generate a draft first."})

        if not user_edits:
            # Just return the current draft
            return json.dumps({
                "draft_id": draft_id,
                "version": draft.get("version", 1),
                "draft_text": draft.get("draft_text", ""),
                "message": "Current draft returned. Provide user_edits to make changes.",
            })

        if regenerate:
            # Ask the AI to cleanly incorporate the edits
            system = (
                "You are a legal document editor. The user has provided edits or instructions "
                "for an existing legal document. Produce a clean, final version of the document "
                "incorporating all changes. Maintain formal legal language and proper structure."
            )
            prompt = (
                f"Original document:\n{draft.get('draft_text', '')}\n\n"
                f"User edits / instructions:\n{user_edits}\n\n"
                "Produce the complete updated document."
            )
            new_text = generate(prompt, system)
        else:
            # Just save the user's edits as-is
            new_text = user_edits

        # Save updated draft
        draft["draft_text"] = new_text
        draft["version"] = draft.get("version", 1) + 1
        draft["updated_at"] = datetime.utcnow().isoformat()
        draft["edit_history"] = draft.get("edit_history", [])
        draft["edit_history"].append({
            "version": draft["version"] - 1,
            "edited_at": datetime.utcnow().isoformat(),
            "user_edits_summary": user_edits[:200],
        })
        _save_draft(draft_id, draft)

        return json.dumps({
            "draft_id": draft_id,
            "version": draft["version"],
            "draft_text": new_text,
            "regenerated": regenerate,
            "message": "Draft updated and saved.",
        })


# ---------------------------------------------------------------------------
# 3. Document Exporter — real file output (TXT always, PDF if reportlab available)
# ---------------------------------------------------------------------------

class DocumentExporterTool(Tool):
    name = "document_exporter"
    description = (
        "Export a draft document to a file. Supports txt (always), "
        "docx (Word format, needs python-docx), and pdf (needs reportlab). "
        "Returns the file path. DOCX is recommended — opens in Word and Google Docs."
    )
    inputs = {
        "draft_id": {
            "type": "string",
            "description": "The draft_id to export.",
        },
        "format": {
            "type": "string",
            "description": "Output format: txt | docx | pdf  (default: docx)",
            "nullable": True,
        },
    }
    output_type = "string"

    def forward(self, draft_id: str, format: str = "docx") -> str:
        draft = _load_draft(draft_id)
        if draft is None:
            return json.dumps({"error": f"Draft '{draft_id}' not found."})

        fmt = (format or "txt").lower().strip()
        draft_text = draft.get("draft_text", "")
        ttype = draft.get("template_type", "document")
        exports_dir = DRAFTS_DIR.parent / "exports"
        exports_dir.mkdir(exist_ok=True)
        filename_base = f"{ttype}_{draft_id}"

        # ── DOCX ──────────────────────────────────────────────────────────
        if fmt == "docx":
            try:
                from docx import Document
                from docx.shared import Pt, Cm, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
                import re

                doc = Document()

                # ── Page setup ────────────────────────────────────────────
                for section in doc.sections:
                    section.top_margin    = Cm(2.54)
                    section.bottom_margin = Cm(2.54)
                    section.left_margin   = Cm(3.0)
                    section.right_margin  = Cm(2.54)

                # ── Style helpers ─────────────────────────────────────────
                def _set_run_font(run, size_pt, bold=False, italic=False, color=None):
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(size_pt)
                    run.bold   = bold
                    run.italic = italic
                    if color:
                        run.font.color.rgb = RGBColor(*color)

                def _add_hrule(doc):
                    """Add a thin horizontal line (paragraph border)."""
                    p = doc.add_paragraph()
                    pPr = p._p.get_or_add_pPr()
                    pBdr = OxmlElement('w:pBdr')
                    bottom = OxmlElement('w:bottom')
                    bottom.set(qn('w:val'), 'single')
                    bottom.set(qn('w:sz'), '6')
                    bottom.set(qn('w:space'), '1')
                    bottom.set(qn('w:color'), '000000')
                    pBdr.append(bottom)
                    pPr.append(pBdr)
                    return p

                def _add_inline_runs(para, text, base_size=11):
                    """
                    Parse inline markdown in `text` and add runs to `para`.
                    Handles: **bold**, *italic*, *(italic)*, `code`
                    """
                    # Pattern: **bold**, *italic*, `code`
                    pattern = re.compile(r'\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`')
                    last = 0
                    for m in pattern.finditer(text):
                        # Plain text before this match
                        if m.start() > last:
                            r = para.add_run(text[last:m.start()])
                            _set_run_font(r, base_size)
                        if m.group(1):   # **bold**
                            r = para.add_run(m.group(1))
                            _set_run_font(r, base_size, bold=True)
                        elif m.group(2): # *italic*
                            r = para.add_run(m.group(2))
                            _set_run_font(r, base_size, italic=True)
                        elif m.group(3): # `code`
                            r = para.add_run(m.group(3))
                            r.font.name = "Courier New"
                            r.font.size = Pt(base_size - 1)
                        last = m.end()
                    # Remaining plain text
                    if last < len(text):
                        r = para.add_run(text[last:])
                        _set_run_font(r, base_size)

                # ── Line-by-line renderer ─────────────────────────────────
                lines = draft_text.split("\n")
                i = 0
                while i < len(lines):
                    line = lines[i]
                    stripped = line.strip()

                    # Skip markdown table separator rows |---|---|
                    if re.match(r'^\|[-|: ]+\|$', stripped):
                        i += 1
                        continue

                    # Horizontal rule ---
                    if re.match(r'^-{3,}$', stripped) or re.match(r'^\*{3,}$', stripped):
                        _add_hrule(doc)
                        i += 1
                        continue

                    # Heading 1: # text
                    if stripped.startswith("# "):
                        text = stripped[2:].strip()
                        # Strip inline markdown from heading
                        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
                        text = re.sub(r'\*(.+?)\*', r'\1', text)
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        r = p.add_run(text.upper())
                        _set_run_font(r, 13, bold=True)
                        i += 1
                        continue

                    # Heading 2: ## text
                    if stripped.startswith("## "):
                        text = stripped[3:].strip()
                        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        r = p.add_run(text.upper())
                        _set_run_font(r, 12, bold=True)
                        # Underline via border
                        pPr = p._p.get_or_add_pPr()
                        pBdr = OxmlElement('w:pBdr')
                        bot = OxmlElement('w:bottom')
                        bot.set(qn('w:val'), 'single')
                        bot.set(qn('w:sz'), '4')
                        bot.set(qn('w:space'), '1')
                        bot.set(qn('w:color'), '000000')
                        pBdr.append(bot)
                        pPr.append(pBdr)
                        i += 1
                        continue

                    # Heading 3: ### text
                    if stripped.startswith("### "):
                        text = stripped[4:].strip()
                        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
                        p = doc.add_paragraph()
                        r = p.add_run(text)
                        _set_run_font(r, 11, bold=True)
                        i += 1
                        continue

                    # Numbered list: "1. text" or "1) text"
                    num_match = re.match(r'^(\d+)[.)]\s+(.+)$', stripped)
                    if num_match:
                        num   = num_match.group(1)
                        text  = num_match.group(2)
                        p = doc.add_paragraph()
                        p.paragraph_format.left_indent  = Cm(0.75)
                        p.paragraph_format.first_line_indent = Cm(-0.75)
                        p.paragraph_format.space_after  = Pt(4)
                        # Number run
                        r = p.add_run(f"{num}. ")
                        _set_run_font(r, 11, bold=False)
                        # Content with inline markdown
                        _add_inline_runs(p, text, base_size=11)
                        i += 1
                        continue

                    # Bullet list: - text or * text
                    if re.match(r'^[-*]\s+', stripped):
                        text = re.sub(r'^[-*]\s+', '', stripped)
                        p = doc.add_paragraph()
                        p.paragraph_format.left_indent       = Cm(0.75)
                        p.paragraph_format.first_line_indent = Cm(-0.5)
                        r_bullet = p.add_run("• ")
                        _set_run_font(r_bullet, 11)
                        _add_inline_runs(p, text, base_size=11)
                        i += 1
                        continue

                    # Table row: | col | col |
                    if stripped.startswith("|") and stripped.endswith("|"):
                        cells = [c.strip() for c in stripped.strip("|").split("|")]
                        cells = [re.sub(r'\*\*(.+?)\*\*', r'\1', c) for c in cells]
                        cells = [re.sub(r'\*(.+?)\*', r'\1', c) for c in cells]
                        cells = [c for c in cells if c]
                        if cells:
                            p = doc.add_paragraph()
                            p.paragraph_format.left_indent = Cm(0.5)
                            for j, cell in enumerate(cells):
                                r = p.add_run(cell)
                                _set_run_font(r, 11, bold=(j == 0))
                                if j < len(cells) - 1:
                                    p.add_run("    ")
                        i += 1
                        continue

                    # Empty line → small spacer
                    if not stripped:
                        p = doc.add_paragraph()
                        p.paragraph_format.space_after = Pt(2)
                        i += 1
                        continue

                    # Normal paragraph with inline markdown
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(4)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    _add_inline_runs(p, stripped, base_size=11)
                    i += 1

                docx_path = exports_dir / f"{filename_base}.docx"
                doc.save(str(docx_path))
                return json.dumps({
                    "draft_id": draft_id,
                    "format": "docx",
                    "file_path": str(docx_path),
                    "message": f"DOCX exported to {docx_path}. Open in Word or Google Docs.",
                })
            except ImportError:
                return json.dumps({"error": "python-docx not installed. Run: pip install python-docx"})
            except Exception as e:
                return json.dumps({"error": f"DOCX generation failed: {e}"})

        # ── PDF ───────────────────────────────────────────────────────────
        if fmt == "pdf":
            try:
                from reportlab.lib.pagesizes import A4
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib.units import cm
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                from reportlab.lib.enums import TA_LEFT

                pdf_path = exports_dir / f"{filename_base}.pdf"
                doc = SimpleDocTemplate(
                    str(pdf_path),
                    pagesize=A4,
                    leftMargin=2.5 * cm,
                    rightMargin=2.5 * cm,
                    topMargin=2.5 * cm,
                    bottomMargin=2.5 * cm,
                )
                styles = getSampleStyleSheet()
                body_style = ParagraphStyle(
                    "Body",
                    parent=styles["Normal"],
                    fontSize=11,
                    leading=16,
                    alignment=TA_LEFT,
                )
                story = []
                for line in draft_text.split("\n"):
                    if line.strip():
                        story.append(Paragraph(line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"), body_style))
                    else:
                        story.append(Spacer(1, 0.3 * cm))
                doc.build(story)
                return json.dumps({
                    "draft_id": draft_id,
                    "format": "pdf",
                    "file_path": str(pdf_path),
                    "message": f"PDF exported to {pdf_path}",
                })
            except ImportError:
                fmt = "txt"  # fall through to txt

        # ── TXT (always works) ────────────────────────────────────────────
        txt_path = exports_dir / f"{filename_base}.txt"
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(draft_text)

        return json.dumps({
            "draft_id": draft_id,
            "format": "txt",
            "file_path": str(txt_path),
            "message": f"Document exported to {txt_path}. Open the file to copy/print.",
        })


# ---------------------------------------------------------------------------
# 4. Authority Finder — AI-powered, real jurisdiction-aware lookup
# ---------------------------------------------------------------------------

class AuthorityFinderTool(Tool):
    name = "authority_finder"
    description = (
        "Find the correct government authority, office address, portal URL, helpline, "
        "and office hours for a given case type and jurisdiction in India."
    )
    inputs = {
        "authority_type": {
            "type": "string",
            "description": (
                "Type of authority needed: police | consumer_court | labour_commission | "
                "rti_officer | cyber_crime | women_commission | human_rights | district_court"
            ),
        },
        "jurisdiction": {
            "type": "string",
            "description": "State and/or city, e.g. 'Maharashtra, Mumbai' or 'Delhi'.",
        },
    }
    output_type = "string"

    def forward(self, authority_type: str, jurisdiction: str) -> str:
        prompt = (
            f"Authority Type: {authority_type}\n"
            f"Jurisdiction: {jurisdiction}\n\n"
            "Provide the specific authority details for this jurisdiction. "
            "Include: authority name, full address, phone/helpline, official portal URL, "
            "office hours, what to bring, and what to say if they refuse to accept the complaint."
        )
        return generate(prompt, SYSTEM_PROMPTS["authority_finder"])


# ---------------------------------------------------------------------------
# 5. Checklist Generator — AI-powered, case-type and step aware
# ---------------------------------------------------------------------------

class ChecklistGeneratorTool(Tool):
    name = "checklist_generator"
    description = (
        "Generate a precise checklist of documents, IDs, fees, and annexures required "
        "for a specific workflow step and case type. Tailored to the actual situation."
    )
    inputs = {
        "workflow_step": {
            "type": "string",
            "description": (
                "The specific step the user is at, e.g. 'filing consumer complaint at district forum', "
                "'submitting RTI application', 'registering FIR at police station'."
            ),
        },
        "case_type": {
            "type": "string",
            "description": "Case type: consumer_complaint | fir | rti | labour_complaint | civil_litigation | other",
        },
        "jurisdiction": {
            "type": "string",
            "description": "State/city for jurisdiction-specific requirements.",
            "nullable": True,
        },
    }
    output_type = "string"

    def forward(self, workflow_step: str, case_type: str, jurisdiction: str = None) -> str:
        prompt = (
            f"Workflow Step: {workflow_step}\n"
            f"Case Type: {case_type}\n"
            f"Jurisdiction: {jurisdiction or 'India (general)'}\n\n"
            "List every document, ID proof, fee, and annexure the user must bring or prepare. "
            "Be specific — include number of copies, certified vs. self-attested, fee amounts, "
            "payment modes, and any forms to fill in advance."
        )
        return generate(prompt, SYSTEM_PROMPTS["checklist_generator"])
