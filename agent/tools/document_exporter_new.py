"""
LaTeX-quality document exporter for NyayaMitr.
Generates professional DOCX and PDF from draft markdown text.
Documents are grouped by session_id in exports/{session_id}/ folders.
"""
import json
import re
from datetime import datetime
from pathlib import Path

EXPORTS_BASE = Path(__file__).resolve().parents[1] / "exports"
EXPORTS_BASE.mkdir(exist_ok=True)


def get_session_exports_dir(session_id: str = None) -> Path:
    """Return exports/{session_id}/ or exports/ if no session."""
    if session_id:
        d = EXPORTS_BASE / session_id
        d.mkdir(parents=True, exist_ok=True)
        return d
    return EXPORTS_BASE


# ─── PDF ──────────────────────────────────────────────────────────────────────

def export_pdf(draft_text: str, out_path: Path, title: str = "Legal Document") -> Path:
    """
    Generate a LaTeX-quality PDF using reportlab.
    Features: proper typography, headers/footers, section styling,
    justified text, table support, page numbers.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm, mm
    from reportlab.lib.colors import HexColor, black, white
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY, TA_RIGHT
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, HRFlowable,
        Table, TableStyle, KeepTogether, PageBreak,
    )
    from reportlab.platypus.flowables import HRFlowable
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    # ── Colours (LaTeX-inspired, professional) ────────────────────────────
    C_BLACK   = HexColor("#1a1a1a")
    C_DARK    = HexColor("#2d2d2d")
    C_ACCENT  = HexColor("#1a3a5c")   # deep navy — like LaTeX article headings
    C_RULE    = HexColor("#2d2d2d")
    C_MUTED   = HexColor("#555555")
    C_LIGHT   = HexColor("#f5f5f5")
    C_BORDER  = HexColor("#cccccc")
    C_GOLD    = HexColor("#8b6914")   # for official stamp / watermark feel

    W, H = A4  # 595.27 x 841.89 pts

    # ── Page template with header/footer ─────────────────────────────────
    doc_title = title
    generated_date = datetime.now().strftime("%d %B %Y")

    def _on_page(canvas, doc):
        canvas.saveState()
        # Top rule
        canvas.setStrokeColor(C_ACCENT)
        canvas.setLineWidth(1.5)
        canvas.line(2.2*cm, H - 1.8*cm, W - 2.2*cm, H - 1.8*cm)
        # Header text
        canvas.setFont("Helvetica-Bold", 7)
        canvas.setFillColor(C_ACCENT)
        canvas.drawString(2.2*cm, H - 1.5*cm, "NYAYAMITR — AI LEGAL ASSISTANCE")
        canvas.setFont("Helvetica", 7)
        canvas.setFillColor(C_MUTED)
        canvas.drawRightString(W - 2.2*cm, H - 1.5*cm, f"Generated: {generated_date}")
        # Bottom rule
        canvas.setStrokeColor(C_ACCENT)
        canvas.setLineWidth(0.5)
        canvas.line(2.2*cm, 1.8*cm, W - 2.2*cm, 1.8*cm)
        # Footer: page number + disclaimer
        canvas.setFont("Helvetica", 7)
        canvas.setFillColor(C_MUTED)
        canvas.drawCentredString(W/2, 1.3*cm, f"Page {doc.page}")
        canvas.setFont("Helvetica-Oblique", 6.5)
        canvas.drawCentredString(W/2, 0.9*cm,
            "This document is for informational purposes only and does not constitute legal advice.")
        canvas.restoreState()

    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=A4,
        leftMargin=2.8*cm,
        rightMargin=2.8*cm,
        topMargin=2.8*cm,
        bottomMargin=2.8*cm,
        title=doc_title,
        author="NyayaMitr AI",
        subject="Legal Document",
    )

    # ── Styles ────────────────────────────────────────────────────────────
    def S(name, **kw):
        return ParagraphStyle(name, **kw)

    base = dict(fontName="Times-Roman", fontSize=11, leading=17,
                textColor=C_BLACK)

    styles = {
        "h1": S("H1", fontName="Helvetica-Bold", fontSize=15, leading=20,
                 textColor=C_ACCENT, alignment=TA_CENTER,
                 spaceBefore=14, spaceAfter=6),
        "h2": S("H2", fontName="Helvetica-Bold", fontSize=12, leading=16,
                 textColor=C_ACCENT, alignment=TA_LEFT,
                 spaceBefore=12, spaceAfter=4),
        "h3": S("H3", fontName="Helvetica-Bold", fontSize=11, leading=15,
                 textColor=C_DARK, alignment=TA_LEFT,
                 spaceBefore=8, spaceAfter=3),
        "body": S("Body", fontName="Times-Roman", fontSize=11, leading=17,
                   textColor=C_BLACK, alignment=TA_JUSTIFY, spaceAfter=4),
        "body_left": S("BodyLeft", fontName="Times-Roman", fontSize=11, leading=17,
                        textColor=C_BLACK, alignment=TA_LEFT, spaceAfter=4),
        "bullet": S("Bullet", fontName="Times-Roman", fontSize=11, leading=17,
                     textColor=C_BLACK, alignment=TA_LEFT,
                     leftIndent=18, firstLineIndent=-12, spaceAfter=3),
        "numbered": S("Numbered", fontName="Times-Roman", fontSize=11, leading=17,
                       textColor=C_BLACK, alignment=TA_JUSTIFY,
                       leftIndent=22, firstLineIndent=-16, spaceAfter=3),
        "label": S("Label", fontName="Helvetica-Bold", fontSize=10,
                    textColor=C_MUTED, leading=14, spaceAfter=2),
        "code": S("Code", fontName="Courier", fontSize=9, leading=13,
                   textColor=C_DARK, backColor=C_LIGHT,
                   leftIndent=12, rightIndent=12, spaceAfter=6),
        "center": S("Center", fontName="Times-Roman", fontSize=11, leading=17,
                     textColor=C_BLACK, alignment=TA_CENTER, spaceAfter=4),
        "right": S("Right", fontName="Times-Roman", fontSize=11, leading=17,
                    textColor=C_BLACK, alignment=TA_RIGHT, spaceAfter=4),
        "small": S("Small", fontName="Times-Roman", fontSize=9,
                    textColor=C_MUTED, leading=13, spaceAfter=2),
    }

    def esc(t):
        """Escape XML special chars for reportlab Paragraph."""
        return (t.replace("&", "&amp;")
                 .replace("<", "&lt;")
                 .replace(">", "&gt;")
                 .replace('"', "&quot;"))

    def inline_md(text):
        """Convert **bold**, *italic*, `code` to reportlab XML tags."""
        text = esc(text)
        text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
        text = re.sub(r'\*(.+?)\*',     r'<i>\1</i>', text)
        text = re.sub(r'`(.+?)`',       r'<font name="Courier" size="9">\1</font>', text)
        return text

    # ── Parse markdown → flowables ────────────────────────────────────────
    story = []
    lines = draft_text.split("\n")
    i = 0

    # Collect table rows
    table_rows = []

    def flush_table():
        nonlocal table_rows
        if not table_rows:
            return
        # Build reportlab Table
        col_count = max(len(r) for r in table_rows)
        # Pad rows
        padded = [r + [""] * (col_count - len(r)) for r in table_rows]
        col_w = (W - 5.6*cm) / col_count
        t = Table(padded, colWidths=[col_w]*col_count, repeatRows=1)
        t.setStyle(TableStyle([
            ("BACKGROUND",   (0,0), (-1,0),  C_ACCENT),
            ("TEXTCOLOR",    (0,0), (-1,0),  white),
            ("FONTNAME",     (0,0), (-1,0),  "Helvetica-Bold"),
            ("FONTSIZE",     (0,0), (-1,0),  9),
            ("FONTNAME",     (0,1), (-1,-1), "Times-Roman"),
            ("FONTSIZE",     (0,1), (-1,-1), 10),
            ("ROWBACKGROUNDS",(0,1),(-1,-1), [white, C_LIGHT]),
            ("GRID",         (0,0), (-1,-1), 0.5, C_BORDER),
            ("VALIGN",       (0,0), (-1,-1), "MIDDLE"),
            ("TOPPADDING",   (0,0), (-1,-1), 5),
            ("BOTTOMPADDING",(0,0), (-1,-1), 5),
            ("LEFTPADDING",  (0,0), (-1,-1), 8),
            ("RIGHTPADDING", (0,0), (-1,-1), 8),
        ]))
        story.append(Spacer(1, 4))
        story.append(t)
        story.append(Spacer(1, 8))
        table_rows.clear()

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Table row
        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            # Skip separator rows |---|
            if all(re.match(r'^[-:]+$', c) for c in cells if c):
                i += 1
                continue
            # Strip bold from header cells
            cells = [re.sub(r'\*\*(.+?)\*\*', r'\1', c) for c in cells]
            table_rows.append(cells)
            i += 1
            continue
        else:
            flush_table()

        # Horizontal rule
        if re.match(r'^-{3,}$', stripped) or re.match(r'^\*{3,}$', stripped):
            story.append(Spacer(1, 4))
            story.append(HRFlowable(width="100%", thickness=0.8,
                                     color=C_RULE, spaceAfter=4))
            i += 1
            continue

        # H1
        if stripped.startswith("# "):
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', stripped[2:].strip())
            story.append(Spacer(1, 8))
            story.append(Paragraph(text.upper(), styles["h1"]))
            story.append(HRFlowable(width="60%", thickness=1.5,
                                     color=C_ACCENT, spaceAfter=6))
            i += 1
            continue

        # H2
        if stripped.startswith("## "):
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', stripped[3:].strip())
            story.append(Spacer(1, 6))
            story.append(Paragraph(text, styles["h2"]))
            story.append(HRFlowable(width="100%", thickness=0.5,
                                     color=C_ACCENT, spaceAfter=3))
            i += 1
            continue

        # H3
        if stripped.startswith("### "):
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', stripped[4:].strip())
            story.append(Paragraph(text, styles["h3"]))
            i += 1
            continue

        # Numbered list
        num_m = re.match(r'^(\d+)[.)]\s+(.+)$', stripped)
        if num_m:
            num, text = num_m.group(1), num_m.group(2)
            p = Paragraph(f"<b>{num}.</b>  {inline_md(text)}", styles["numbered"])
            story.append(p)
            i += 1
            continue

        # Bullet
        if re.match(r'^[-*•]\s+', stripped):
            text = re.sub(r'^[-*•]\s+', '', stripped)
            p = Paragraph(f"•  {inline_md(text)}", styles["bullet"])
            story.append(p)
            i += 1
            continue

        # Empty line
        if not stripped:
            story.append(Spacer(1, 5))
            i += 1
            continue

        # Normal paragraph
        story.append(Paragraph(inline_md(stripped), styles["body"]))
        i += 1

    flush_table()

    doc.build(story, onFirstPage=_on_page, onLaterPages=_on_page)
    return out_path


# ─── DOCX ─────────────────────────────────────────────────────────────────────

def export_docx(draft_text: str, out_path: Path, title: str = "Legal Document") -> Path:
    """
    Generate a LaTeX-quality DOCX using python-docx.
    Features: Times New Roman body, proper heading hierarchy,
    header/footer with page numbers, table styling, justified text.
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_SECTION
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    import lxml.etree as etree

    # ── Colours ───────────────────────────────────────────────────────────
    NAVY   = RGBColor(0x1a, 0x3a, 0x5c)
    BLACK  = RGBColor(0x1a, 0x1a, 0x1a)
    MUTED  = RGBColor(0x55, 0x55, 0x55)
    GOLD   = RGBColor(0x8b, 0x69, 0x14)
    LIGHT  = RGBColor(0xf5, 0xf5, 0xf5)

    doc = Document()

    # ── Page setup ────────────────────────────────────────────────────────
    for section in doc.sections:
        section.page_width  = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin    = Cm(2.8)
        section.bottom_margin = Cm(2.8)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.8)

    # ── Header ────────────────────────────────────────────────────────────
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Left: brand name
    r1 = hp.add_run("NYAYAMITR — AI LEGAL ASSISTANCE")
    r1.font.name = "Calibri"
    r1.font.size = Pt(7)
    r1.font.bold = True
    r1.font.color.rgb = NAVY
    # Tab to right side
    hp.add_run("\t")
    r2 = hp.add_run(f"Generated: {datetime.now().strftime('%d %B %Y')}")
    r2.font.name = "Calibri"
    r2.font.size = Pt(7)
    r2.font.color.rgb = MUTED
    # Bottom border on header paragraph
    pPr = hp._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '12')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '1a3a5c')
    pBdr.append(bot)
    pPr.append(pBdr)
    # Tab stop at right margin
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), '9072')  # ~16cm in twips
    tabs.append(tab)
    pPr.append(tabs)

    # ── Footer with page number ───────────────────────────────────────────
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Top border
    fpPr = fp._p.get_or_add_pPr()
    fpBdr = OxmlElement('w:pBdr')
    ftop = OxmlElement('w:top')
    ftop.set(qn('w:val'), 'single')
    ftop.set(qn('w:sz'), '4')
    ftop.set(qn('w:space'), '1')
    ftop.set(qn('w:color'), '1a3a5c')
    fpBdr.append(ftop)
    fpPr.append(fpBdr)
    # Page number field
    r_pg = fp.add_run()
    r_pg.font.name = "Calibri"
    r_pg.font.size = Pt(8)
    r_pg.font.color.rgb = MUTED
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = ' PAGE '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    r_pg._r.append(fldChar1)
    r_pg._r.append(instrText)
    r_pg._r.append(fldChar2)
    # Disclaimer
    fp.add_run("\n")
    r_disc = fp.add_run(
        "This document is for informational purposes only and does not constitute legal advice."
    )
    r_disc.font.name = "Calibri"
    r_disc.font.size = Pt(6.5)
    r_disc.font.italic = True
    r_disc.font.color.rgb = MUTED

    # ── Style helpers ─────────────────────────────────────────────────────
    def _font(run, size, bold=False, italic=False, color=None, name="Times New Roman"):
        run.font.name = name
        run.font.size = Pt(size)
        run.bold   = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color

    def _para_border_bottom(para, color_hex="1a3a5c", sz="6"):
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single')
        bot.set(qn('w:sz'), sz)
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), color_hex)
        pBdr.append(bot)
        pPr.append(pBdr)

    def _add_inline(para, text, size=11):
        """Add inline markdown runs to a paragraph."""
        pattern = re.compile(r'\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`')
        last = 0
        for m in pattern.finditer(text):
            if m.start() > last:
                r = para.add_run(text[last:m.start()])
                _font(r, size)
            if m.group(1):
                r = para.add_run(m.group(1))
                _font(r, size, bold=True)
            elif m.group(2):
                r = para.add_run(m.group(2))
                _font(r, size, italic=True)
            elif m.group(3):
                r = para.add_run(m.group(3))
                r.font.name = "Courier New"
                r.font.size = Pt(size - 1)
            last = m.end()
        if last < len(text):
            r = para.add_run(text[last:])
            _font(r, size)

    def _add_hrule(color_hex="2d2d2d", sz="6"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single')
        bot.set(qn('w:sz'), sz)
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), color_hex)
        pBdr.append(bot)
        pPr.append(pBdr)

    # ── Parse markdown → DOCX elements ───────────────────────────────────
    lines = draft_text.split("\n")
    i = 0
    table_rows = []

    def flush_table():
        nonlocal table_rows
        if not table_rows:
            return
        col_count = max(len(r) for r in table_rows)
        padded = [r + [""] * (col_count - len(r)) for r in table_rows]
        t = doc.add_table(rows=len(padded), cols=col_count)
        t.style = "Table Grid"
        for ri, row_data in enumerate(padded):
            for ci, cell_text in enumerate(row_data):
                cell = t.cell(ri, ci)
                cell.text = ""
                p = cell.paragraphs[0]
                r = p.add_run(re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text))
                _font(r, 10, bold=(ri == 0))
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                # Header row shading
                if ri == 0:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), '1a3a5c')
                    tcPr.append(shd)
                    r.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
                elif ri % 2 == 0:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), 'f5f5f5')
                    tcPr.append(shd)
        doc.add_paragraph()
        table_rows.clear()

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Table row
        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            if all(re.match(r'^[-:]+$', c) for c in cells if c):
                i += 1
                continue
            cells = [re.sub(r'\*\*(.+?)\*\*', r'\1', c) for c in cells]
            table_rows.append(cells)
            i += 1
            continue
        else:
            flush_table()

        # Horizontal rule
        if re.match(r'^-{3,}$', stripped) or re.match(r'^\*{3,}$', stripped):
            _add_hrule()
            i += 1
            continue

        # H1
        if stripped.startswith("# "):
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', stripped[2:].strip())
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after  = Pt(6)
            r = p.add_run(text.upper())
            _font(r, 14, bold=True, color=NAVY, name="Calibri")
            _para_border_bottom(p, "1a3a5c", "12")
            i += 1
            continue

        # H2
        if stripped.startswith("## "):
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', stripped[3:].strip())
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(4)
            r = p.add_run(text)
            _font(r, 12, bold=True, color=NAVY, name="Calibri")
            _para_border_bottom(p, "1a3a5c", "4")
            i += 1
            continue

        # H3
        if stripped.startswith("### "):
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', stripped[4:].strip())
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(3)
            r = p.add_run(text)
            _font(r, 11, bold=True, color=BLACK)
            i += 1
            continue

        # Numbered list
        num_m = re.match(r'^(\d+)[.)]\s+(.+)$', stripped)
        if num_m:
            num, text = num_m.group(1), num_m.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent       = Cm(0.8)
            p.paragraph_format.first_line_indent = Cm(-0.8)
            p.paragraph_format.space_after        = Pt(3)
            r_num = p.add_run(f"{num}.  ")
            _font(r_num, 11, bold=True)
            _add_inline(p, text, 11)
            i += 1
            continue

        # Bullet
        if re.match(r'^[-*•]\s+', stripped):
            text = re.sub(r'^[-*•]\s+', '', stripped)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent       = Cm(0.8)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            p.paragraph_format.space_after        = Pt(3)
            r_b = p.add_run("•  ")
            _font(r_b, 11, color=NAVY)
            _add_inline(p, text, 11)
            i += 1
            continue

        # Empty line
        if not stripped:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(5)
        _add_inline(p, stripped, 11)
        i += 1

    flush_table()

    doc.save(str(out_path))
    return out_path
